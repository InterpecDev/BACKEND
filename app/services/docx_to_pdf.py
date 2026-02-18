import os
from datetime import datetime
from docx import Document
from docx2pdf import convert  # Windows (usa Word)

def _replace_in_paragraph(p, mapping: dict[str, str]):
    # Reemplazo robusto: Word suele partir texto en runs
    full = "".join(r.text for r in p.runs)
    changed = False
    for k, v in mapping.items():
        if k in full:
            full = full.replace(k, v)
            changed = True
    if changed:
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = full
        else:
            p.add_run(full)

def fill_template_docx(template_path: str, out_docx_path: str, mapping: dict[str, str]):
    doc = Document(template_path)

    # Párrafos
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    # Tablas
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

    os.makedirs(os.path.dirname(out_docx_path), exist_ok=True)
    doc.save(out_docx_path)

def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    out_dir = os.path.dirname(pdf_path)
    os.makedirs(out_dir, exist_ok=True)

    convert(docx_path, out_dir)

    generated = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if not os.path.exists(generated):
        raise RuntimeError("No se pudo generar el PDF (Word/docx2pdf).")

    # Renombrar al nombre esperado
    if generated != pdf_path:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        os.rename(generated, pdf_path)

def build_pdf_from_template(
    template_path: str,
    storage_dir: str,
    base_name: str,
    mapping: dict[str, str],
) -> tuple[str, str]:
    stamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    out_dir = os.path.join(storage_dir, "dictamenes")

    out_docx = os.path.join(out_dir, f"{base_name}_{stamp}.docx")
    out_pdf  = os.path.join(out_dir, f"{base_name}_{stamp}.pdf")

    fill_template_docx(template_path, out_docx, mapping)
    convert_docx_to_pdf(out_docx, out_pdf)
    return out_docx, out_pdf